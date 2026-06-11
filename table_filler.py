"""Fill the 3-tomadas medication table template with parsed medications.

The template has a landscape table with 4 columns:
  Col 0: Medication name + dose
  Col 1: Morning (sun icon)
  Col 2: Lunch (plate icon)
  Col 3: Night (moon icon)

Row 0: Merged title header (patient name)
Row 1: Icon headers
Rows 2+: Data rows for medications
"""

import copy
import re
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models import ParsedMedication


# ---------------------------------------------------------------------------
# Posology / timing → (morning, lunch, night) mapping
# ---------------------------------------------------------------------------

def _parse_posology_code(code: str) -> tuple[str, str, str]:
    """Parse a posology code like '1-0-1' into (morning, lunch, night) strings.

    Values > 1 are kept as-is (e.g., '2-0-1' → ('2', '', '1')).
    Values of 0 become empty strings.
    Values of 1 become 'X' (checkmark placeholder).
    """
    parts = code.split("-")
    if len(parts) != 3:
        return ("", "", "")

    result = []
    for p in parts:
        try:
            val = int(p)
        except ValueError:
            result.append(p)
            continue
        if val == 0:
            result.append("")
        else:
            result.append(str(val))
    return (result[0], result[1], result[2])


def _timing_to_slots(timing: Optional[str]) -> tuple[str, str, str]:
    """Map a timing string to (morning, lunch, night) slots."""
    if not timing:
        return ("", "", "")

    t = timing.lower().strip()
    morning = lunch = night = ""

    # Check for compound timing: "cedo e à noite", "manhã e à tarde"
    if " e " in t:
        parts = t.split(" e ", 1)
        for part in parts:
            part = part.strip()
            if _is_morning(part):
                morning = "1"
            elif _is_lunch(part):
                lunch = "1"
            elif _is_night(part):
                night = "1"
    else:
        if _is_morning(t):
            morning = "1"
        elif _is_lunch(t):
            lunch = "1"
        elif _is_night(t):
            night = "1"

    return (morning, lunch, night)


def _is_morning(t: str) -> bool:
    return bool(re.search(r"cedo|manh[aã]|jejum", t))


def _is_lunch(t: str) -> bool:
    return bool(re.search(r"almo[cç]o|tarde|meio[- ]?dia", t))


def _is_night(t: str) -> bool:
    return bool(re.search(r"noite|jantar|deitar", t))


def _frequency_to_slots(frequency: Optional[str]) -> tuple[str, str, str]:
    """Infer time slots from frequency when no posology/timing is available."""
    if not frequency:
        return ("", "", "")

    f = frequency.lower()

    # 8/8h or de 8 em 8 horas → 3x/day
    if re.search(r"8\s*/\s*8|de\s+8\s+em\s+8", f):
        return ("1", "1", "1")

    # 12/12h or de 12 em 12 horas → 2x/day (morning + night)
    if re.search(r"12\s*/\s*12|de\s+12\s+em\s+12", f):
        return ("1", "", "1")

    # 6/6h → 4x/day — mark all three
    if re.search(r"6\s*/\s*6|de\s+6\s+em\s+6", f):
        return ("1", "1", "1")

    # "3 vezes" or "3x" → all three
    m = re.search(r"(\d+)\s*(?:vezes|x)", f)
    if m:
        n = int(m.group(1))
        if n >= 3:
            return ("1", "1", "1")
        elif n == 2:
            return ("1", "", "1")
        elif n == 1:
            return ("1", "", "")

    # "1 vez ao dia" → morning by default
    if re.search(r"1\s*vez\s*(?:ao|por|no)\s*dia", f):
        return ("1", "", "")

    return ("", "", "")


def get_time_slots(med: ParsedMedication) -> tuple[str, str, str]:
    """Determine (morning, lunch, night) content for a parsed medication.

    Priority: posology_code > timing > frequency.
    Returns the quantity_per_dose in each active slot when available.
    """
    # 1. Posology code is most explicit
    if med.posology_code:
        slots = _parse_posology_code(med.posology_code)
        # Replace numeric count with quantity_per_dose if available (e.g. "1cp")
        if med.quantity_per_dose:
            qty = med.quantity_per_dose
            slots = tuple(qty if s and s != "" else "" for s in slots)
        return slots

    # 2. Timing
    if med.timing:
        slots = _timing_to_slots(med.timing)
        if any(slots):
            if med.quantity_per_dose:
                qty = med.quantity_per_dose
                slots = tuple(qty if s else "" for s in slots)
            return slots

    # 3. Frequency
    if med.frequency:
        slots = _frequency_to_slots(med.frequency)
        if any(slots):
            if med.quantity_per_dose:
                qty = med.quantity_per_dose
                slots = tuple(qty if s else "" for s in slots)
            return slots

    # 4. No timing info at all — put X in morning as default for 1x/day meds
    return ("1", "", "")


# ---------------------------------------------------------------------------
# Build medication display name
# ---------------------------------------------------------------------------

def _build_med_display(med: ParsedMedication) -> str:
    """Build the text for column 0: medication name + dose."""
    parts = [med.medication_name]
    if med.dose:
        parts.append(med.dose)
    return " ".join(parts)


# ---------------------------------------------------------------------------
# DOCX table filling
# ---------------------------------------------------------------------------

def _clone_row(table, template_row_idx: int):
    """Clone a row from the table (deep copy of XML) and append it."""
    tbl = table._tbl
    template_tr = tbl.tr_lst[template_row_idx]
    new_tr = copy.deepcopy(template_tr)
    tbl.append(new_tr)
    return new_tr


def _set_cell_text(cell, text: str, font_name: str = "Times New Roman",
                   font_size: int = 20):
    """Set text of a cell, preserving alignment but replacing content."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""

    # Use first paragraph
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)


def fill_table(
    template_path: str,
    medications: list[ParsedMedication],
    patient_name: str = "",
) -> BytesIO:
    """Fill the 3-tomadas template with parsed medications.

    Returns a BytesIO with the generated .docx.
    """
    doc = Document(template_path)
    table = doc.tables[0]

    # Row 0 is the merged title header — set patient name if provided
    if patient_name:
        _set_cell_text(table.rows[0].cells[0], patient_name,
                       font_size=20)

    # Row 1 has the icons — leave untouched
    # Rows 2..8 are the 7 existing empty data rows

    first_data_row = 2
    existing_data_rows = len(table.rows) - first_data_row  # 7

    # Filter out suspended and TARV medications
    active_meds = [
        m for m in medications
        if m.suspension.value == "active" and not m.is_tarv
    ]

    for i, med in enumerate(active_meds):
        if i < existing_data_rows:
            # Fill existing row
            row = table.rows[first_data_row + i]
        else:
            # Clone the last data row template and append
            _clone_row(table, first_data_row + existing_data_rows - 1)
            row = table.rows[-1]

        display_name = _build_med_display(med)
        morning, lunch, night = get_time_slots(med)

        _set_cell_text(row.cells[0], display_name)
        _set_cell_text(row.cells[1], morning)
        _set_cell_text(row.cells[2], lunch)
        _set_cell_text(row.cells[3], night)

    # Remove unused empty data rows
    total_needed = len(active_meds)
    total_data_available = len(table.rows) - first_data_row
    if total_needed < total_data_available:
        tbl_xml = table._tbl
        rows_to_remove = []
        for j in range(first_data_row + total_needed, len(table.rows)):
            rows_to_remove.append(tbl_xml.tr_lst[j])
        for tr in rows_to_remove:
            tbl_xml.remove(tr)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
